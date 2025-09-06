from fastapi import APIRouter, Depends, HTTPException, Response
from sqlalchemy.orm import Session
from app import models, database
from fastapi.responses import StreamingResponse
import io
from docx import Document
from weasyprint import HTML

router = APIRouter()

def get_db():
    db = database.SessionLocal()
    try:
        yield db
    finally:
        db.close()

def render_resume_docx(user: models.User) -> bytes:
    doc = Document()
    doc.add_heading(user.name, 0)
    doc.add_paragraph(user.email)
    doc.add_heading("Experience", level=1)
    for exp in user.experiences:
        p = doc.add_paragraph()
        p.add_run(f"{exp.role.name} at {exp.company.name} ({exp.start_year}-{exp.end_year or 'Present'})\n").bold = True
        if exp.description:
            doc.add_paragraph(exp.description)
    doc.add_heading("Skills", level=1)
    for skill in user.skills:
        doc.add_paragraph(f"{skill.name} ({skill.level or ''})")
    doc.add_heading("Education", level=1)
    for edu in user.educations:
        doc.add_paragraph(f"{edu.degree} in {edu.field or ''}, {edu.school.name} ({edu.start_year}-{edu.end_year})")
    doc.add_heading("Interests", level=1)
    for interest in user.interests:
        doc.add_paragraph(interest.name)
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    return f.read()

def render_resume_html(user: models.User) -> str:
    html = f"""
    <html><body>
    <h1>{user.name}</h1>
    <p>{user.email}</p>
    <h2>Experience</h2>
    <ul>
    """
    for exp in user.experiences:
        html += f"<li><b>{exp.role.name} at {exp.company.name} ({exp.start_year}-{exp.end_year or 'Present'})</b><br>"
        if exp.description:
            html += f"{exp.description}" 
        html += "</li>"
    html += "</ul><h2>Skills</h2><ul>"
    for skill in user.skills:
        html += f"<li>{skill.name} ({skill.level or ''})</li>"
    html += "</ul><h2>Education</h2><ul>"
    for edu in user.educations:
        html += f"<li>{edu.degree} in {edu.field or ''}, {edu.school.name} ({edu.start_year}-{edu.end_year})</li>"
    html += "</ul><h2>Interests</h2><ul>"
    for interest in user.interests:
        html += f"<li>{interest.name}</li>"
    html += "</ul></body></html>"
    return html

@router.get("/profile/{user_id}/resume.docx")
def export_resume_docx(user_id: int, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    docx_bytes = render_resume_docx(user)
    return Response(docx_bytes, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=resume_{user_id}.docx"})

@router.get("/profile/{user_id}/resume.pdf")
def export_resume_pdf(user_id: int, db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.id == user_id).first()
    if not user:
        raise HTTPException(status_code=404, detail="User not found")
    html = render_resume_html(user)
    pdf_bytes = HTML(string=html).write_pdf()
    return Response(pdf_bytes, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=resume_{user_id}.pdf"})
